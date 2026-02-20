from __future__ import annotations

"""
ODIE report generator
--------------------
This module generates a DOCX report from a list of DisasterEvent objects.

Design goals:
- Keep ODIE usable even if report dependencies are missing (lazy imports).
- Choose charts that match the *current result set*.
  Example: If the user filtered to ONE country and ONE disaster type,
  then "Top countries" and "Top disaster types" charts are not informative,
  so we switch to subtype/time/severity charts automatically.
- Add beginner-friendly comments so the code is easy to explain in a DSA report.
"""

from dataclasses import dataclass, field
from typing import List, Optional, Sequence, Tuple
import os
import tempfile
from collections import Counter
import math

from .models import DisasterEvent


# -----------------------------
# Configuration / citation types
# -----------------------------

@dataclass
class DatasetCitation:
    """Minimal dataset citation metadata for the DOCX report."""
    database_name: str = "Emergency Events Database (EM-DAT)"
    institutional_author: str = "UCLouvain / CRED"
    location: str = "Brussels, Belgium"
    access_date_iso: str = "2026-01-30"
    website: str = "https://www.emdat.be"
    file_name: Optional[str] = None
    file_note: Optional[str] = "Custom export (Excel) generated from EM-DAT."


@dataclass
class ReportConfig:
    """High-level knobs to control how the report is written."""
    title: str = "ODIE Analytical Report"
    subtitle: str = "Offline Disaster Intelligence Engine (CLI)"
    dataset_name: str = "EM-DAT Excel export"
    citation: DatasetCitation = field(default_factory=DatasetCitation)

    # How many categories to show in bar charts / tables
    top_n: int = 10

    # How many rows to show in preview tables
    max_rows_preview: int = 15

    # Optional: list of CLI commands used to create the current result set
    command_log: Optional[List[str]] = None


# -----------------------------
# Helpers for clean numeric plots
# -----------------------------

def _safe_floats(values: Sequence[Optional[float]]) -> List[float]:
    """Convert a list with Nones to clean floats (skip NaN/inf)."""
    out: List[float] = []
    for v in values:
        if v is None:
            continue
        try:
            fv = float(v)
        except Exception:
            continue
        if math.isfinite(fv):
            out.append(fv)
    return out


def _choose_bins(n: int) -> int:
    """Simple bin heuristic (keeps charts readable for small samples)."""
    if n <= 20:
        return 10
    if n <= 100:
        return 15
    return 30


# -----------------------------
# Main entry point used by CLI
# -----------------------------

def generate_docx_report(
    events: Sequence[DisasterEvent],
    out_path: str,
    *,
    config: Optional[ReportConfig] = None,
    scope_label: str = "Current Result Set",
) -> str:
    """
    Generate a DOCX report + charts for a list of events.

    IMPORTANT:
    - This does NOT modify your Excel file.
    - ODIE loads data into memory and reports on that in-memory selection.
    """
    config = config or ReportConfig()

    # Lazy imports: only required when "report" is used.
    # This way ODIE can run without python-docx/matplotlib until needed.
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError as e:
        raise ImportError(
            "Missing dependency: python-docx.\n"
            "Install it with: python -m pip install python-docx"
        ) from e

    try:
        import matplotlib.pyplot as plt
        import numpy as np
    except ImportError as e:
        raise ImportError(
            "Missing dependency: matplotlib (and numpy).\n"
            "Install with: python -m pip install matplotlib numpy"
        ) from e

    if not events:
        raise ValueError("No events to report on (result set is empty).")

    # -----------------------------
    # 1) Compute stats + category counts
    # -----------------------------
    years = [e.start_year for e in events if e.start_year is not None]
    deaths = _safe_floats([e.total_deaths for e in events])
    affected = _safe_floats([e.total_affected for e in events])
    damage = _safe_floats([e.total_damage_adj_usd for e in events])

    countries = [e.country for e in events if e.country]
    types = [e.disaster_type for e in events if e.disaster_type]
    subtypes = [e.disaster_subtype for e in events if e.disaster_subtype]

    c_country = Counter(countries)
    c_type = Counter(types)
    c_subtype = Counter(subtypes)

    unique_countries = len(c_country)
    unique_types = len(c_type)

    year_min = min(years) if years else None
    year_max = max(years) if years else None

    # -----------------------------
    # 2) Create charts (SMART selection)
    # -----------------------------
    tmpdir = tempfile.mkdtemp(prefix="odie_report_")
    # Each chart is: (title, file_path, why_this_chart)
    chart_paths: List[Tuple[str, str, str]] = []

    def _save(filename: str) -> str:
        path = os.path.join(tmpdir, filename)
        plt.tight_layout()
        plt.savefig(path, dpi=200)
        plt.close()
        return path

    def _bar(title: str, labels: List[str], values: List[int], why: str, filename: str) -> None:
        plt.figure()
        plt.bar(labels, values)
        plt.xticks(rotation=45, ha="right")
        plt.title(title)
        plt.ylabel("Count")
        chart_paths.append((title, _save(filename), why))

    def _scatter(title: str, x: List[float], y: List[float], xlabel: str, ylabel: str, why: str, filename: str) -> None:
        plt.figure()
        plt.scatter(x, y)
        plt.title(title)
        plt.xlabel(xlabel)
        plt.ylabel(ylabel)
        chart_paths.append((title, _save(filename), why))

    def _hist_log10_alternating(title: str, data: List[float], xlabel: str, why: str, filename: str) -> None:
        """
        Histogram with:
        - visible bin boundaries (edgecolor)
        - alternating bar colors (to visually separate ranges)
        - log10 transform to handle extreme outliers
        """
        if not data:
            return
        x = np.log10(np.array(data) + 1.0)
        plt.figure()
        counts, bins, patches = plt.hist(
            x,
            bins=_choose_bins(len(x)),
            edgecolor="black",
            linewidth=0.8,
        )
        for i, p in enumerate(patches):
            # User requested alternating colors
            p.set_facecolor("C0" if i % 2 == 0 else "C1")
        plt.title(title)
        plt.xlabel(xlabel)
        plt.ylabel("Count")
        chart_paths.append((title, _save(filename), why))

    # Year distribution: histogram if enough unique years, else decade bar
    if years:
        if len(set(years)) >= 15:
            plt.figure()
            plt.hist([float(y) for y in years], bins=_choose_bins(len(years)), edgecolor="black", linewidth=0.8)
            plt.title(f"Event distribution by Start Year ({scope_label})")
            plt.xlabel("Start Year")
            plt.ylabel("Count")
            chart_paths.append((
                f"Event distribution by Start Year ({scope_label})",
                _save("hist_years.png"),
                "Histogram is suitable for showing how events are spread across years."
            ))
        else:
            decades = sorted(set((y // 10) * 10 for y in years))
            decade_counts = [sum(1 for y in years if (y // 10) * 10 == d) for d in decades]
            _bar(
                f"Event count by Decade ({scope_label})",
                [str(d) for d in decades],
                decade_counts,
                "Bar chart is clearer than a histogram when years are few or clustered.",
                "bar_decade.png"
            )

    # Category charts:
    # - If multiple countries/types -> show top countries/types
    # - If only one country/type -> show subtype distribution instead
    if unique_countries > 1:
        top_countries = c_country.most_common(config.top_n)
        _bar(
            f"Top {config.top_n} Countries by Number of Events ({scope_label})",
            [k for k, _ in top_countries],
            [v for _, v in top_countries],
            "Bar charts are ideal for comparing category counts (countries).",
            "top_countries.png"
        )
    else:
        top_sub = c_subtype.most_common(config.top_n)
        if top_sub:
            _bar(
                f"Top {config.top_n} Disaster Subtypes by Number of Events ({scope_label})",
                [k for k, _ in top_sub],
                [v for _, v in top_sub],
                "When the subset is a single country/type, subtype comparison becomes informative.",
                "top_subtypes.png"
            )

    if unique_types > 1:
        top_types = c_type.most_common(config.top_n)
        _bar(
            f"Top {config.top_n} Disaster Types by Number of Events ({scope_label})",
            [k for k, _ in top_types],
            [v for _, v in top_types],
            "Bar charts are ideal for comparing category counts (disaster types).",
            "top_types.png"
        )

    # Severity over time: scatter plots
    if years and deaths:
        xs = [float(e.start_year) for e in events if e.start_year is not None and e.total_deaths is not None]
        ys = [float(e.total_deaths) for e in events if e.start_year is not None and e.total_deaths is not None]
        if xs and ys:
            _scatter(
                f"Total Deaths vs Start Year ({scope_label})",
                xs, ys,
                "Start Year", "Total Deaths",
                "Scatter plot shows how severity changes over time without forcing bins.",
                "scatter_deaths.png"
            )

    if years and affected:
        xs = [float(e.start_year) for e in events if e.start_year is not None and e.total_affected is not None]
        ys = [float(e.total_affected) for e in events if e.start_year is not None and e.total_affected is not None]
        if xs and ys:
            _scatter(
                f"Total Affected vs Start Year ({scope_label})",
                xs, ys,
                "Start Year", "Total Affected",
                "Scatter plot highlights unusually large affected counts and trends across years.",
                "scatter_affected.png"
            )

    # Severity distribution: histogram (log-scale, alternating colors)
    if deaths:
        _hist_log10_alternating(
            f"Distribution of Total Deaths (log10 scale) ({scope_label})",
            deaths,
            "log10(Total Deaths + 1)",
            "Deaths are heavy-tailed. Log scale + bin boundaries improves readability.",
            "hist_deaths_log.png"
        )
    elif affected:
        _hist_log10_alternating(
            f"Distribution of Total Affected (log10 scale) ({scope_label})",
            affected,
            "log10(Total Affected + 1)",
            "Affected counts can be heavy-tailed. Log scale improves readability.",
            "hist_affected_log.png"
        )
    elif damage:
        _hist_log10_alternating(
            f"Distribution of Adjusted Damage (log10 scale) ({scope_label})",
            damage,
            "log10(Adjusted Damage + 1)",
            "Damage values often have extreme outliers. Log scale improves readability.",
            "hist_damage_log.png"
        )

    # -----------------------------
    # 3) Build DOCX report
    # -----------------------------
    doc = Document()

    # Set a simple readable default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    style.font.size = Pt(11)

    def _center_title(text: str, size: int, bold: bool = False, italic: bool = False) -> None:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _kv(key: str, value: str) -> None:
        p = doc.add_paragraph()
        r = p.add_run(f"{key}: ")
        r.bold = True
        p.add_run(value)

    _center_title(config.title, 22, bold=True)
    _center_title(config.subtitle, 12, italic=True)

    doc.add_paragraph("")
    _kv("Dataset", config.dataset_name)
    _kv("Scope", scope_label)
    _kv("Total records in scope", str(len(events)))
    if year_min is not None and year_max is not None:
        _kv("Year range (Start Year)", f"{year_min} to {year_max}")

    # Dataset citation section
    doc.add_paragraph("")
    doc.add_heading("Dataset citation", level=1)
    cit = config.citation
    if cit.file_name:
        doc.add_paragraph(f"Data file used: {cit.file_name}")
    if cit.file_note:
        doc.add_paragraph(f"File note: {cit.file_note}")
    doc.add_paragraph("Suggested citation (database):")
    doc.add_paragraph(
        f"{cit.institutional_author} (accessed {cit.access_date_iso}). "
        f"{cit.database_name}. {cit.location}. {cit.website}."
    )

    # Command log (optional) for reproducibility
    if config.command_log:
        doc.add_paragraph("")
        doc.add_heading("Command log (reproducibility)", level=1)
        doc.add_paragraph("These ODIE commands produced this result set:")
        for line in config.command_log:
            doc.add_paragraph(line, style="List Bullet")

    # Data dictionary
    doc.add_paragraph("")
    doc.add_heading("Columns used (data dictionary)", level=1)
    doc.add_paragraph("This report uses the following ODIE fields derived from dataset columns:")
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "ODIE field"
    t.rows[0].cells[1].text = "Meaning"
    for k, v in [
        ("dis_no", "Unique disaster identifier"),
        ("country", "Country name"),
        ("disaster_type", "Major disaster category"),
        ("disaster_subtype", "Subtype within the category"),
        ("start_year", "Event start year"),
        ("total_deaths", "Total deaths"),
        ("total_affected", "Total affected"),
        ("total_damage_adj_usd", "Adjusted damage (US$)"),
    ]:
        row = t.add_row().cells
        row[0].text = k
        row[1].text = v

    # Data completeness summary (missingness)
    doc.add_paragraph("")
    doc.add_heading("Data completeness", level=1)
    doc.add_paragraph("Availability of numeric fields in this scope:")
    t2 = doc.add_table(rows=1, cols=3)
    t2.rows[0].cells[0].text = "Metric"
    t2.rows[0].cells[1].text = "Available"
    t2.rows[0].cells[2].text = "Missing"

    def _add_missing_row(name: str, getter) -> None:
        vals = [getter(e) for e in events]
        missing = sum(1 for v in vals if v is None)
        available = len(events) - missing
        row = t2.add_row().cells
        row[0].text = name
        row[1].text = str(available)
        row[2].text = str(missing)

    _add_missing_row("Total Deaths", lambda e: e.total_deaths)
    _add_missing_row("Total Affected", lambda e: e.total_affected)
    _add_missing_row("Adjusted Damage (US$)", lambda e: e.total_damage_adj_usd)

    # Visualizations
    doc.add_paragraph("")
    doc.add_heading("Visualizations", level=1)
    for title, path, why in chart_paths:
        doc.add_paragraph(title)
        doc.add_picture(path, width=Inches(6.5))
        doc.add_paragraph("Why this graph is suitable: " + why)
        doc.add_paragraph("")

    # Key event tables (top-k by deaths/damage)
    doc.add_paragraph("")
    doc.add_heading("Key events tables", level=1)

    top_by_deaths = sorted(
        [e for e in events if e.total_deaths is not None],
        key=lambda e: float(e.total_deaths),
        reverse=True
    )[:10]
    if top_by_deaths:
        doc.add_paragraph("Top 10 events by Total Deaths (within scope)")
        t3 = doc.add_table(rows=1, cols=7)
        h = t3.rows[0].cells
        h[0].text = "DisNo"
        h[1].text = "Country"
        h[2].text = "Type"
        h[3].text = "Subtype"
        h[4].text = "Year"
        h[5].text = "Deaths"
        h[6].text = "Adj. Damage (US$)"
        for e in top_by_deaths:
            r = t3.add_row().cells
            r[0].text = e.dis_no
            r[1].text = e.country
            r[2].text = e.disaster_type
            r[3].text = e.disaster_subtype
            r[4].text = str(e.start_year if e.start_year is not None else "")
            r[5].text = str(int(e.total_deaths)) if e.total_deaths is not None else ""
            r[6].text = f"{int(e.total_damage_adj_usd):,}" if e.total_damage_adj_usd is not None else ""

    top_by_damage = sorted(
        [e for e in events if e.total_damage_adj_usd is not None],
        key=lambda e: float(e.total_damage_adj_usd),
        reverse=True
    )[:10]
    if top_by_damage:
        doc.add_paragraph("")
        doc.add_paragraph("Top 10 events by Adjusted Damage (US$) (within scope)")
        t4 = doc.add_table(rows=1, cols=6)
        h = t4.rows[0].cells
        h[0].text = "DisNo"
        h[1].text = "Year"
        h[2].text = "Subtype"
        h[3].text = "Adj. Damage (US$)"
        h[4].text = "Deaths"
        h[5].text = "Total Affected"
        for e in top_by_damage:
            r = t4.add_row().cells
            r[0].text = e.dis_no
            r[1].text = str(e.start_year if e.start_year is not None else "")
            r[2].text = e.disaster_subtype
            r[3].text = f"{int(e.total_damage_adj_usd):,}" if e.total_damage_adj_usd is not None else ""
            r[4].text = str(int(e.total_deaths)) if e.total_deaths is not None else ""
            r[5].text = str(int(e.total_affected)) if e.total_affected is not None else ""

    # A small preview table (first N records)
    doc.add_paragraph("")
    doc.add_heading("Preview of first few records", level=1)
    preview = list(events)[:config.max_rows_preview]
    t5 = doc.add_table(rows=1, cols=6)
    h = t5.rows[0].cells
    h[0].text = "DisNo"
    h[1].text = "Country"
    h[2].text = "Type"
    h[3].text = "Subtype"
    h[4].text = "Year"
    h[5].text = "Deaths"
    for e in preview:
        r = t5.add_row().cells
        r[0].text = e.dis_no
        r[1].text = e.country
        r[2].text = e.disaster_type
        r[3].text = e.disaster_subtype
        r[4].text = str(e.start_year if e.start_year is not None else "")
        r[5].text = str(int(e.total_deaths)) if e.total_deaths is not None else ""
    doc.add_heading("Notes", level=1)
    doc.add_paragraph(
        "ODIE does not modify the original Excel dataset. "
        "Filters and queries change only the selected record IDs in memory. "
        "This report is generated from that selected subset."
    )

    
    # -----------------------------
    # Reproducibility footer (professional report touch)
    # -----------------------------
    doc.add_paragraph("")
    doc.add_heading("Reproducibility footer", level=1)

    # ODIE version (from the package)
    try:
        from . import __version__ as odie_version
    except Exception:
        odie_version = "unknown"

    # Timestamp (when the report was generated)
    from datetime import datetime as _dt
    generated_at = _dt.now().isoformat(timespec="seconds")

    doc.add_paragraph(f"ODIE version: {odie_version}")
    doc.add_paragraph(f"Report generated at: {generated_at}")
    doc.add_paragraph(f"Records in scope: {len(events)}")

    if config.citation.file_name:
        doc.add_paragraph(f"Dataset file: {config.citation.file_name}")

    # Command log (again) at the end so it is easy to find
    if config.command_log:
        doc.add_paragraph("Commands used (log):")
        for line in config.command_log:
            doc.add_paragraph(line, style="List Bullet")

    # Short algorithmic notes (ties report to DSA syllabus)
    doc.add_paragraph("")
    doc.add_paragraph("Algorithmic notes (DSA):")
    for note in [
        "Filtering uses precomputed indices (value -> sorted list of IDs) and list intersection.",
        "Year-range filtering uses binary search (bisect) over sorted years.",
        "Sorting uses comparison-based sorting (merge sort / quick sort utilities).",
        "Top-k queries use a heap (priority queue).",
        "Advanced 'where' filtering parses an expression into an AST (tree) and evaluates it."
    ]:
        doc.add_paragraph(note, style="List Bullet")

    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)
    return out_path
